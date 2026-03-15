import json
import tkinter as tk
import os
from tkinter import ttk
from docx import Document
from datetime import datetime


def load_last():
    if os.path.exists("last_data.json"):
        with open("last_data.json", "r", encoding="utf8") as f:
            return json.load(f)
    return {}


def save_last(data):
    with open("last_data.json", "w", encoding="utf8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


def load_json(file):
    with open(file, "r", encoding="utf8") as f:
        return json.load(f)


works = load_json("works.json")
employees = load_json("employees.json")["employees"]

employee_names = [e["name"] for e in employees]


def refresh_employee_lists():
    global employee_names
    employee_names = [e["name"] for e in employees]

    for cb in [executor_cb, leader_cb, rb_cb, chief_cb, issued_cb, accepted_cb]:
        cb["values"] = employee_names


def get_post(name):
    for e in employees:
        if e["name"] == name:
            return e["post"]
    return ""


def replace_text(doc, data):

    BOLD_FIELDS = [
        "{{jobs}}",
        "{{material}}",
        "{{protect}}",
        "{{rboblast}}",
        "{{controls}}"
    ]

    for p in doc.paragraphs:

        for key, value in data.items():

            if key in p.text:

                before, after = p.text.split(key, 1)

                p.clear()

                p.add_run(before)

                run = p.add_run(value)
                run.underline = True

                if key in BOLD_FIELDS:
                    run.bold = True

                p.add_run(after)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in data.items():
                        if key in p.text:
                            p.text = p.text.replace(key, value)


def fill_workers_table(doc, workers, date):

    for table in doc.tables:

        for row in table.rows:

            if "{{worker}}" in row.cells[1].text:

                template = row
                table._tbl.remove(template._tr)

                for i, w in enumerate(workers):

                    row = table.add_row()

                    row.cells[0].text = str(i + 1)
                    row.cells[1].text = w[0]
                    row.cells[2].text = w[1]
                    row.cells[3].text = date

                return


def generate():

    executor = executor_cb.get()
    leader = leader_cb.get()
    rb = rb_cb.get()
    chief = chief_cb.get()
    issued = issued_cb.get()
    accepted = accepted_cb.get()

    team = team_cb.get()
    date = date_entry.get()

    dateDDMM = date[:5]
    dateYYYY = "20" + date[-2:]

    workers = []

    for w in worker_rows:

        name = w.get()

        if name:
            workers.append((name, get_post(name)))

    data = {

        "{{executor}}": executor,
        "{{team}}": team,

        "{{leader}}": leader,
        "{{leadePost}}": get_post(leader),

        "{{rb}}": rb,
        "{{rbPost}}": get_post(rb),

        "{{chief}}": chief,
        "{{chiefPost}}": get_post(chief),

        "{{issued}}": issued,
        "{{issuedPost}}": get_post(issued),

        "{{accepted}}": accepted,
        "{{acceptedPost}}": get_post(accepted),

        "{{jobs}}": jobs_cb.get(),
        "{{material}}": material_cb.get(),
        "{{protect}}": protect_cb.get(),
        "{{rboblast}}": rboblast_cb.get(),
        "{{controls}}": controls_cb.get(),

        "{{data}}": date,
        "{{dataDDMM}}": dateDDMM,
        "{{dataYYYY}}": dateYYYY
    }

    doc = Document("template.docx")

    replace_text(doc, data)

    fill_workers_table(doc, workers, date)

    doc.save("narad_ready.docx")

    save_last({
        "executor": executor,
        "leader": leader,
        "rb": rb,
        "chief": chief,
        "issued": issued,
        "accepted": accepted,
        "team": team,
        "date": date,
        "jobs": jobs_cb.get(),
        "material": material_cb.get(),
        "protect": protect_cb.get(),
        "rboblast": rboblast_cb.get(),
        "controls": controls_cb.get(),
        "workers": [w.get() for w in worker_rows if w.get()]
    })

def add_employee_window():

    win = tk.Toplevel(root)
    win.title("Добавить сотрудника")

    tk.Label(win, text="ФИО").grid(row=0, column=0)

    name_entry = tk.Entry(win, width=30)
    name_entry.grid(row=0, column=1)

    tk.Label(win, text="Должность").grid(row=1, column=0)

    post_entry = tk.Entry(win, width=30)
    post_entry.grid(row=1, column=1)

    def save_employee():

        name = name_entry.get()
        post = post_entry.get()

        if not name:
            return

        employees.append({"name": name, "post": post})

        with open("employees.json", "w", encoding="utf8") as f:
            json.dump({"employees": employees}, f, ensure_ascii=False, indent=4)

        refresh_employee_lists()

        win.destroy()

    tk.Button(win, text="Сохранить", command=save_employee).grid(row=2, column=1)
def add_work_window():

    win = tk.Toplevel(root)
    win.title("Добавить вариант")

    tk.Label(win, text="Тип").grid(row=0, column=0)

    type_cb = ttk.Combobox(win, values=[
        "jobs",
        "material",
        "protect",
        "rboblast",
        "controls"
    ])
    type_cb.grid(row=0, column=1)

    tk.Label(win, text="Текст").grid(row=1, column=0)

    text_entry = tk.Entry(win, width=50)
    text_entry.grid(row=1, column=1)

    def save_work():

        t = type_cb.get()
        text = text_entry.get()

        if not t or not text:
            return

        works[t].append(text)

        with open("works.json", "w", encoding="utf8") as f:
            json.dump(works, f, ensure_ascii=False, indent=4)

        # обновляем списки в главном окне
        jobs_cb["values"] = works["jobs"]
        material_cb["values"] = works["material"]
        protect_cb["values"] = works["protect"]
        rboblast_cb["values"] = works["rboblast"]
        controls_cb["values"] = works["controls"]

        win.destroy()

    tk.Button(win, text="Сохранить", command=save_work).grid(row=2, column=1)

def manage_data():

    win = tk.Toplevel(root)
    win.title("Управление данными")

    notebook = ttk.Notebook(win)
    notebook.pack(fill="both", expand=True)

    emp_frame = tk.Frame(notebook)
    work_frame = tk.Frame(notebook)

    notebook.add(emp_frame, text="Сотрудники")
    notebook.add(work_frame, text="Работы")

    emp_list = tk.Listbox(emp_frame, width=50)
    emp_list.pack()

    def refresh_emp():

        emp_list.delete(0, tk.END)

        for e in employees:
            emp_list.insert(tk.END, f"{e['name']} — {e['post']}")

    refresh_emp()

    def delete_employee():

        sel = emp_list.curselection()

        if not sel:
            return

        index = sel[0]

        employees.pop(index)

        with open("employees.json", "w", encoding="utf8") as f:
            json.dump({"employees": employees}, f, ensure_ascii=False, indent=4)

        refresh_emp()
        refresh_employee_lists()

    def edit_employee():

        sel = emp_list.curselection()

        if not sel:
            return

        index = sel[0]

        e = employees[index]

        edit = tk.Toplevel(win)

        tk.Label(edit, text="ФИО").grid(row=0)
        name = tk.Entry(edit)
        name.insert(0, e["name"])
        name.grid(row=0, column=1)

        tk.Label(edit, text="Должность").grid(row=1)
        post = tk.Entry(edit)
        post.insert(0, e["post"])
        post.grid(row=1, column=1)

        def save():

            employees[index]["name"] = name.get()
            employees[index]["post"] = post.get()

            with open("employees.json", "w", encoding="utf8") as f:
                json.dump({"employees": employees}, f, ensure_ascii=False, indent=4)

            refresh_emp()
            refresh_employee_lists()

            edit.destroy()

        tk.Button(edit, text="Сохранить", command=save).grid(row=2, column=1)

    tk.Button(emp_frame, text="Изменить", command=edit_employee).pack()

    tk.Button(emp_frame, text="Удалить", command=delete_employee).pack()

    work_list = tk.Listbox(work_frame, width=70)
    work_list.pack()

    work_names = {
        "jobs": "Работы",
        "material": "Материалы",
        "protect": "СИЗ",
        "rboblast": "Рад. обстановка",
        "controls": "Контроль"
    }

    def refresh_work():

        work_list.delete(0, tk.END)

        for k, vals in works.items():

            for v in vals:
                name = work_names.get(k, k)

                work_list.insert(tk.END, f"{name} : {v}")

    refresh_work()

    def delete_work():

        sel = work_list.curselection()

        if not sel:
            return

        line = work_list.get(sel)

        name, text = line.split(" : ", 1)

        reverse_names = {
            "Работы": "jobs",
            "Материалы": "material",
            "СИЗ": "protect",
            "Рад. обстановка": "rboblast",
            "Контроль": "controls"
        }

        key = reverse_names[name]

        works[key].remove(text)

        with open("works.json", "w", encoding="utf8") as f:
            json.dump(works, f, ensure_ascii=False, indent=4)

        refresh_work()

    tk.Button(work_frame, text="Удалить", command=delete_work).pack()


root = tk.Tk()

last = load_last()

root.title("Наряд-допуск")

frame = tk.Frame(root)
frame.pack(padx=10, pady=10)


def add_combo(label, values, row):

    tk.Label(frame, text=label).grid(row=row, column=0, sticky="w")

    cb = ttk.Combobox(frame, values=values, width=35)
    cb.grid(row=row, column=1)

    return cb


def add_entry(label, row):

    tk.Label(frame, text=label).grid(row=row, column=0, sticky="w")

    entry = tk.Entry(frame, width=38)
    entry.grid(row=row, column=1)

    return entry


executor_cb = add_combo("Ответственный исполнитель", employee_names, 0)
leader_cb = add_combo("Руководитель работ", employee_names, 1)
rb_cb = add_combo("Ответственный РБ", employee_names, 2)
chief_cb = add_combo("Ответственный за работы", employee_names, 3)

issued_cb = add_combo("Наряд выдал", employee_names, 4)
accepted_cb = add_combo("Наряд принял", employee_names, 5)

tk.Label(frame, text="Бригада").grid(row=6, column=0, sticky="w")

team_cb = ttk.Combobox(frame, values=[1,2,3,4,5,6,7,8,9,10], width=35)
team_cb.grid(row=6, column=1)
team_cb.set(last.get("team", 1))

today = datetime.now().strftime("%d.%m.%y")

date_entry = add_entry("Дата", 7)
date_entry.insert(0, last.get("date", today))

jobs_cb = add_combo("Работы", works["jobs"], 8)
material_cb = add_combo("Материалы", works["material"], 9)
protect_cb = add_combo("СИЗ", works["protect"], 10)
rboblast_cb = add_combo("Рад. обстановка", works["rboblast"], 11)
controls_cb = add_combo("Контроль", works["controls"], 12)

jobs_cb.set(last.get("jobs", works["jobs"][0]))
material_cb.set(last.get("material", works["material"][0]))
protect_cb.set(last.get("protect", works["protect"][0]))
rboblast_cb.set(last.get("rboblast", works["rboblast"][0]))
controls_cb.set(last.get("controls", works["controls"][0]))

executor_cb.set(last.get("executor", ""))
leader_cb.set(last.get("leader", ""))
rb_cb.set(last.get("rb", ""))
chief_cb.set(last.get("chief", ""))
issued_cb.set(last.get("issued", ""))
accepted_cb.set(last.get("accepted", ""))

tk.Label(frame, text="Исполнители").grid(row=13)

worker_rows = []
worker_widgets = []
saved_workers = last.get("workers", [])


def add_worker(name=""):

    r = len(worker_rows) + 14

    name_cb = ttk.Combobox(frame, values=employee_names, width=35)
    name_cb.grid(row=r, column=0)

    if name:
        name_cb.set(name)

    worker_rows.append(name_cb)
    worker_widgets.append(name_cb)


def remove_worker():

    if worker_widgets:

        widget = worker_widgets.pop()
        widget.destroy()

        worker_rows.pop()


tk.Button(frame, text="Добавить исполнителя", command=add_worker).grid(row=14, column=2)
tk.Button(frame, text="Удалить исполнителя", command=remove_worker).grid(row=15, column=2)

for w in saved_workers:
    add_worker(w)
tk.Button(frame, text="Добавить сотрудника", command=add_employee_window).grid(row=39, column=0)

tk.Button(frame, text="Добавить вариант работ", command=add_work_window).grid(row=39, column=2)

tk.Button(frame, text="Управление данными", command=manage_data).grid(row=39, column=1)

tk.Button(frame, text="Создать наряд", command=generate).grid(row=40, column=1, pady=10)

root.mainloop()